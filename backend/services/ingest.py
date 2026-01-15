"""Document ingestion service with PDF and DOCX extraction."""

import logging
import asyncio
import io
from typing import List, Optional
from concurrent.futures import ThreadPoolExecutor
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from models.document import Document, DocumentStatus
from models.chunk import Chunk
from core.storage import storage_service
from core.config import settings

logger = logging.getLogger(__name__)

# Thread pool for CPU-bound extraction operations
_executor = ThreadPoolExecutor(max_workers=2)


# ============================================================================
# PDF EXTRACTION
# ============================================================================

def _extract_text_from_pdf_sync(file_bytes: bytes) -> str:
    """
    Synchronous PDF text extraction using pdfplumber.
    
    Args:
        file_bytes: PDF file content as bytes
        
    Returns:
        Extracted text content
    """
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber not installed. Run: pip install pdfplumber")
        return "[PDF extraction requires pdfplumber. Please install it with: pip install pdfplumber]"
    
    try:
        extracted_text = []
        total_chars = 0
        
        # Use BytesIO to avoid writing to disk
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            total_pages = len(pdf.pages)
            max_pages = min(total_pages, settings.PDF_MAX_PAGES)
            
            logger.info(f"Processing PDF: {total_pages} pages (processing first {max_pages})")
            
            for page_num, page in enumerate(pdf.pages[:max_pages]):
                try:
                    # Extract text from page
                    page_text = page.extract_text()
                    
                    if page_text:
                        extracted_text.append(page_text)
                        total_chars += len(page_text)
                    
                    # Log progress for large PDFs
                    if (page_num + 1) % 10 == 0:
                        logger.debug(f"Processed {page_num + 1}/{max_pages} pages")
                        
                except Exception as page_error:
                    logger.warning(f"Error extracting page {page_num + 1}: {str(page_error)}")
                    continue
            
            if total_pages > max_pages:
                extracted_text.append(f"\n\n[Note: PDF truncated. Showing first {max_pages} of {total_pages} pages]")
        
        result = "\n\n".join(extracted_text)
        
        if not result.strip():
            logger.warning("No text extracted from PDF (may be scanned/image-based)")
            return "[No text could be extracted from this PDF. It may be a scanned document or image-based PDF.]"
        
        logger.info(f"PDF extraction complete: {len(extracted_text)} pages, {total_chars} characters")
        return result
        
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}", exc_info=True)
        return f"[PDF extraction failed: {str(e)}]"


async def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Async wrapper for PDF text extraction with timeout protection.
    
    Args:
        file_bytes: PDF file content as bytes
        
    Returns:
        Extracted text content
    """
    if not settings.ENABLE_PDF_EXTRACTION:
        return "[PDF extraction is disabled]"
    
    try:
        loop = asyncio.get_event_loop()
        result = await asyncio.wait_for(
            loop.run_in_executor(_executor, _extract_text_from_pdf_sync, file_bytes),
            timeout=settings.PDF_TIMEOUT_SECONDS
        )
        return result
    except asyncio.TimeoutError:
        logger.error(f"PDF extraction timed out after {settings.PDF_TIMEOUT_SECONDS} seconds")
        return f"[PDF extraction timed out after {settings.PDF_TIMEOUT_SECONDS} seconds]"
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        return f"[PDF extraction error: {str(e)}]"


# ============================================================================
# DOCX EXTRACTION
# ============================================================================

def _extract_text_from_docx_sync(file_bytes: bytes) -> str:
    """
    Synchronous DOCX text extraction using python-docx.
    
    Args:
        file_bytes: DOCX file content as bytes
        
    Returns:
        Extracted text content
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return "[DOCX extraction requires python-docx. Please install it with: pip install python-docx]"
    
    try:
        extracted_text = []
        paragraph_count = 0
        table_count = 0
        total_chars = 0
        
        # Use BytesIO to avoid writing to disk
        doc = DocxDocument(io.BytesIO(file_bytes))
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text.append(para.text)
                paragraph_count += 1
                total_chars += len(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))
            
            if table_text:
                extracted_text.append("\n" + "\n".join(table_text) + "\n")
                table_count += 1
                total_chars += sum(len(t) for t in table_text)
        
        result = "\n\n".join(extracted_text)
        
        if not result.strip():
            logger.warning("No text extracted from DOCX")
            return "[No text could be extracted from this DOCX document]"
        
        logger.info(f"DOCX extraction complete: {paragraph_count} paragraphs, {table_count} tables, {total_chars} characters")
        return result
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}", exc_info=True)
        return f"[DOCX extraction failed: {str(e)}]"


async def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Async wrapper for DOCX text extraction with timeout protection.
    
    Args:
        file_bytes: DOCX file content as bytes
        
    Returns:
        Extracted text content
    """
    if not settings.ENABLE_DOCX_EXTRACTION:
        return "[DOCX extraction is disabled]"
    
    try:
        loop = asyncio.get_event_loop()
        result = await asyncio.wait_for(
            loop.run_in_executor(_executor, _extract_text_from_docx_sync, file_bytes),
            timeout=settings.DOCX_TIMEOUT_SECONDS
        )
        return result
    except asyncio.TimeoutError:
        logger.error(f"DOCX extraction timed out after {settings.DOCX_TIMEOUT_SECONDS} seconds")
        return f"[DOCX extraction timed out after {settings.DOCX_TIMEOUT_SECONDS} seconds]"
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        return f"[DOCX extraction error: {str(e)}]"


# ============================================================================
# MAIN TEXT EXTRACTION
# ============================================================================

async def extract_text(file_bytes: bytes, content_type: str) -> str:
    """
    Extract text from file bytes based on content type.
    
    Args:
        file_bytes: File content as bytes
        content_type: MIME type of the file
        
    Returns:
        Extracted text content
    """
    content_type_lower = content_type.lower()
    
    # Plain text and markdown
    if content_type_lower in ["text/plain", "text/markdown"]:
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            return file_bytes.decode('utf-8', errors='replace')
    
    # PDF files
    elif content_type_lower == "application/pdf":
        logger.info("Extracting text from PDF")
        return await extract_text_from_pdf(file_bytes)
    
    # DOCX files
    elif "wordprocessingml" in content_type_lower or content_type_lower.endswith("msword"):
        logger.info("Extracting text from DOCX")
        return await extract_text_from_docx(file_bytes)
    
    # Fallback: try to decode as text
    else:
        try:
            return file_bytes.decode('utf-8', errors='replace')
        except Exception as e:
            logger.error(f"Failed to extract text from {content_type}: {str(e)}")
            return ""


# ============================================================================
# CHUNKING
# ============================================================================

async def chunk_text(text: str, chunk_size: int, overlap: int) -> List[str]:
    """
    Chunk text using sliding window approach.
    
    Args:
        text: Text to chunk
        chunk_size: Size of each chunk in characters
        overlap: Number of characters to overlap between chunks
        
    Returns:
        List of text chunks
    """
    if not text or chunk_size <= 0:
        return []
    
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]
        
        if chunk.strip():  # Only add non-empty chunks
            chunks.append(chunk)
        
        # Move start position forward by (chunk_size - overlap)
        start += chunk_size - overlap
        
        # Prevent infinite loop if overlap >= chunk_size
        if overlap >= chunk_size:
            break
    
    return chunks


# ============================================================================
# EMBEDDING GENERATION
# ============================================================================

async def generate_chunk_embeddings(chunks: List[Chunk]) -> None:
    """
    Generate embeddings for chunks if enabled.
    
    Args:
        chunks: List of Chunk objects to generate embeddings for
    """
    if not settings.ENABLE_EMBEDDINGS:
        logger.debug("Embeddings disabled, skipping embedding generation")
        return
    
    try:
        from services.embeddings import embedding_service
        
        # Extract text content from chunks
        texts = [chunk.content for chunk in chunks]
        
        # Generate embeddings in batch
        embeddings = await embedding_service.generate_embeddings_batch(texts)
        
        # Store embeddings in chunks
        for chunk, embedding in zip(chunks, embeddings):
            if embedding is not None:
                chunk.embedding = embedding_service.embedding_to_json(embedding)
        
        logger.info(f"Generated embeddings for {len(chunks)} chunks")
        
    except ImportError:
        logger.warning("Embedding service not available, skipping embedding generation")
    except Exception as e:
        logger.error(f"Failed to generate embeddings: {str(e)}")
        # Don't fail ingestion if embeddings fail


# ============================================================================
# DOCUMENT INGESTION
# ============================================================================

async def ingest_document(document_id: str, db: AsyncSession) -> None:
    """
    Ingest a document: download, extract text, chunk, and store chunks.
    
    Args:
        document_id: UUID of the document to ingest
        db: Database session
    """
    try:
        # Load document from database
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        
        if not document:
            logger.error(f"Document {document_id} not found")
            return
        
        # Update status to processing
        document.status = DocumentStatus.PROCESSING
        await db.commit()
        
        # Download file from storage
        logger.info(f"Downloading document {document_id} from storage")
        file_bytes = await storage_service.download_file(document.s3_key)
        
        # Extract text
        logger.info(f"Extracting text from document {document_id} (type: {document.content_type})")
        text = await extract_text(file_bytes, document.content_type)
        
        if not text or text.strip() == "":
            raise ValueError("No text content extracted from document")
        
        # Check if extraction failed (error message)
        if text.startswith("[") and text.endswith("]"):
            logger.warning(f"Document extraction returned message: {text}")
            # Store the message but continue - it's valid content
        
        # Chunk text
        logger.info(f"Chunking text for document {document_id}")
        text_chunks = await chunk_text(
            text,
            settings.CHUNK_SIZE,
            settings.CHUNK_OVERLAP
        )
        
        if not text_chunks:
            raise ValueError("No chunks created from document text")
        
        # Create chunk records
        logger.info(f"Creating {len(text_chunks)} chunks for document {document_id}")
        chunk_objects = []
        for idx, chunk_content in enumerate(text_chunks):
            chunk = Chunk(
                document_id=document_id,
                content=chunk_content,
                chunk_index=idx,
                token_count=len(chunk_content.split())  # Simple word count as token estimate
            )
            chunk_objects.append(chunk)
        
        # Generate embeddings if enabled
        await generate_chunk_embeddings(chunk_objects)
        
        # Save chunks to database
        db.add_all(chunk_objects)
        
        # Update document status to completed
        document.status = DocumentStatus.COMPLETED
        
        # Store extraction stats in metadata
        if not document.metadata_json:
            document.metadata_json = {}
        document.metadata_json['chunks_created'] = len(chunk_objects)
        document.metadata_json['total_characters'] = len(text)
        document.metadata_json['embeddings_generated'] = settings.ENABLE_EMBEDDINGS
        
        await db.commit()
        
        logger.info(f"Successfully ingested document {document_id} with {len(text_chunks)} chunks")
        
    except Exception as e:
        logger.error(f"Failed to ingest document {document_id}: {str(e)}", exc_info=True)
        
        # Update document status to failed
        try:
            result = await db.execute(
                select(Document).where(Document.id == document_id)
            )
            document = result.scalar_one_or_none()
            if document:
                document.status = DocumentStatus.FAILED
                # Store error message in metadata
                if not document.metadata_json:
                    document.metadata_json = {}
                document.metadata_json['error'] = str(e)
                await db.commit()
        except Exception as commit_error:
            logger.error(f"Failed to update document status: {str(commit_error)}")
        
        raise
